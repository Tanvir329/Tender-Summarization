import streamlit as st
import cohere
import PyPDF2
from docx import Document
from io import BytesIO
from docx.shared import Pt
import gspread
from google.oauth2.service_account import Credentials
from datetime import datetime
import pytz
import re
import json
import requests
from dateutil import parser

# Initialize Cohere client
co = cohere.ClientV2(api_key="okYrKAw1OPZoMnOSCR6rUVO2cbSulB4gCmuo04UY")  # Replace with your key

#initialising data variable
data = None

def log_to_google_sheet(filename, file_data, extracted_text):
    scopes = ["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
    credentials = Credentials.from_service_account_info(
        st.secrets["google_sheets"],
        scopes=scopes
    )
    client = gspread.authorize(credentials)
    sheet = client.open("TenderUsageLogs").sheet1

    # Add headers if sheet is empty
    if sheet.row_count == 0 or sheet.cell(1, 1).value != "Timestamp":
        sheet.insert_row(["Timestamp", "Filename", "File Size (KB)", "Text Length", "User IP"], 1)

    # Current time in IST
    ist = pytz.timezone("Asia/Kolkata")
    timestamp = datetime.now(ist).strftime("%Y-%m-%d %H:%M:%S")

    # User IP (fallback to N/A)
    user_ip = st.request.remote_addr if hasattr(st, 'request') else "N/A"

    # File size in KB
    file_data.seek(0, 2)
    file_size_kb = round(file_data.tell() / 1024, 2)
    file_data.seek(0)

    # Extracted text length
    text_length = len(extracted_text.strip())

    # Append the log
    sheet.append_row([timestamp, filename, f"{file_size_kb} KB", text_length, user_ip])

def extract_tender_info(text):
    TenderName = None
    TenderType = None
    StartDate = None
    EndDate = None

    # A single regex to match dd[-/.]mm[-/.]yyyy or Month DD, YYYY
    date_regex = (
        r"(\d{1,2}[-/.]\d{2}[-/.]\d{4}"
        r"|\b(?:January|February|March|April|May|June|July|August|September|October|November|December)"
        r"\s+\d{1,2},\s+\d{4})"
    )

    lines = text.splitlines()
    for idx, raw in enumerate(lines):
        ln = raw.strip()

        # 1) Inline Tender Name match (highest priority)
        m_name_inline = re.search(r"Tender\s+Name\*{0,2}\s*[:\-]\s*(.+)", ln, re.IGNORECASE)
        if m_name_inline:
            TenderName = m_name_inline.group(1).strip()
            # once we get it inline, no need to check heading
            continue

        # 2) Heading Tender Name (if inline wasn't present)
        if TenderName is None and re.match(r"^####\s*\*{0,2}Tender\s+Name\*{0,2}", ln, re.IGNORECASE):
            # next non-empty, non-heading, non-bullet line
            for nxt in lines[idx+1:]:
                cand = nxt.strip()
                if cand and not cand.startswith("####") and not cand.startswith("-"):
                    TenderName = cand
                    break
        
        # List of common labels that can refer to Tender Type
        tender_type_keywords = [
            r"Tender\s+Type", r"Type\s+of\s+Tender", r"Tender\s+Category"
        ]

        # Join them into a single regex pattern using alternation
        label_pattern = "|".join(tender_type_keywords)

        # Updated regex to capture different formats (inline or bullet)
        tender_type_regex = re.compile(
            rf"(?:^|\n|\r|[\*\-•])\s*(?:{label_pattern})\*{{0,2}}\s*[:\-–]\s*(.+)",
            re.IGNORECASE
        )
                # Sample line (you can loop this over all lines in your text)
        m_type = tender_type_regex.search(ln)
        if m_type:
            TenderType = m_type.group(1).strip()

        # 4) Start Date
        if re.search(r"Start\s+Date", ln, re.IGNORECASE):
            m = re.search(date_regex, ln)
            if m:
                raw_date = m.group(1)
                try:
                    parsed_date = parser.parse(raw_date)
                    StartDate = parsed_date.strftime('%Y-%m-%d')  # Format to YYYY-MM-DD
                except Exception as e:
                    print(f"Date parsing failed: {e}")

        # 5) End Date
        if re.search(r"End\s+Date", ln, re.IGNORECASE):
            m = re.search(date_regex, ln)
            if m:
                raw_date = m.group(1)
                try:
                    parsed_date = parser.parse(raw_date)
                    EndDate = parsed_date.strftime('%Y-%m-%d')  # Format to YYYY-MM-DD
                except Exception as e:
                    print(f"Date parsing failed: {e}")
    print(TenderType)
    return {
        "TenderName": TenderName,
        "TenderType": TenderType,
        "StartDate": StartDate,
        "EndDate": EndDate,
    }

def extract_text_from_pdf(pdf_file):
    reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
    return text

def generate_table_word(summary_text):
    # Split lines and extract the first heading (any level)
    lines = summary_text.splitlines()
    heading = next((l.strip().lstrip('#').strip() for l in lines if l.strip().startswith('#')), 'Table')

    # Parse keys and their bullet values
    data = []
    key = None
    values = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.lstrip('#').strip()
        
        # Check if this is a heading/key line
        if re.match(r'^\*\*.*\*\*$', stripped):  # **Key**
            if key:
                # Add the previous key and its values to our data
                data.append((key, values))
            
            key = stripped.strip('*').strip()
            values = []
            
            # Look ahead to check if the next line is a heading or content
            next_index = i + 1
            # Skip empty lines
            while next_index < len(lines) and not lines[next_index].strip():
                next_index += 1
                
            # If we have content but no more lines, or next line is not a heading
            if next_index >= len(lines) or not re.match(r'^#+\s+\*\*.*\*\*$', lines[next_index].strip()):
                # Check if next line is not a bullet point but has content
                if next_index < len(lines) and not lines[next_index].strip().startswith('-') and lines[next_index].strip():
                    # Add the entire paragraph as a value
                    values.append(lines[next_index].strip())
      
        elif key and (stripped.startswith('-') or re.match(r'^\d+\.', stripped)):
            # Remove leading -, 1., 2. etc.
            cleaned = re.sub(r'^[-\d.]+\s*', '', stripped)
            # Prevent duplicates
            if cleaned not in values:
                values.append(cleaned)
            
        i += 1
    
    # Don't forget the last key-value pair
    if key:
        data.append((key, values))

    # Create Word document and add heading
    doc = Document()
    title_para = doc.add_heading(level=1)
    run_title = title_para.add_run(heading)
    run_title.bold = True
    run_title.font.size = Pt(16)

    doc.add_paragraph()  # spacing

    # Build table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Description'
    for cell in hdr_cells:
        for p in cell.paragraphs:
            p.alignment = 1  # Center
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)

    # Populate table
    for key, vals in data:
        row_cells = table.add_row().cells
        cell_key = row_cells[0]
        p_key = cell_key.paragraphs[0]
        run_key = p_key.add_run(key)
        run_key.bold = True
        run_key.font.size = Pt(11)

        cell_val = row_cells[1]
        if vals:
            for v in vals:
                p = cell_val.add_paragraph(style='List Bullet')
                parts = re.split(r'(\*\*[^*]+\*\*)', v)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part.strip('*'))
                        run.bold = True
                    else:
                        p.add_run(part)
        else:
            cell_val.text = ''

    # Save to BytesIO
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def stream_summary_from_cohere(text):
    global data
    x = ""
    prompt = (
        """You are an expert in analyzing and summarizing government and institutional tender documents.

            Summarize the following tender document by extracting and presenting all important and relevant information that may be present. Only include the sections that are explicitly mentioned or applicable in the document. **Do not include sections that are not present, not mentioned, or not relevant to the specific tender type.**

            Extract details under the following categories **only if available**:
            - Tender Name
            - Tender Reference Number and ID  
            - Name of the Issuing Organization or Authority  
            - Tender Fee (amount, mode of payment)  
            - EMD (Earnest Money Deposit) Details (amount, mode of payment)  
            - Estimated Tender Value or Project Cost  
            - Pre-bid Meeting Dates, Venue, and Registration/Link  
            - Tender Meeting Dates and Venues (if different from Pre-bid)  
            - Scope of Work  
            - Modules or Work Packages  
            - Workforce Requirements (specify onsite manpower and training manpower, if any)  
            - Human Resource Details  
            - Technical and Financial Eligibility Criteria  
            - Technical and Financial Marking/Scoring Criteria  
            - Performance Security Requirements  
            - Implementation Timeline and Phases (Turnaround Time or TAT)  
            - Contract Duration/Period  
            - Project Location(s)  
            - Existing IHMS or Software Application Details (if mentioned)  
            - Payment Terms and Schedule  
            - Submission Method (Online, Physical, or Hybrid)  
            - Selection Methodology (e.g., QCBS, L1)  
            - Cloud Service Provider (CSP) Details (if applicable)  
            - Hardware Details (especially for hospital/lab tenders—CT/MRI/X-ray/Pathology equipment)  
            - Technical Specifications  
            - Radiology/Pathology Scope (if applicable)  
            - Checklists (All the documents required, if provided)  
            - Declarations, Undertakings, and Affidavits  
            - Consortium/ Joint Venture
            - OEM (Original Equipment Manufacturer) Document Requirements  
            - Penalty Clauses and Bidder Obligations  
            - Financial Bid Structure  
            - Viability Gap Funding (VGF)  
            - Special Purpose Vehicle (SPV) clauses  
            - Land Border Sharing Clause  
            - Mode of Payments for Tender Fee, EMD, and Other Charges  
            - Contact Details of the Tender Issuer (email, phone, address)

            Present the summary in a clean, organized format using clear headings or bullet points. Again, include **only the sections that are actually present in the document** and dont say not mentioned in the document, instead skip that section.
            
            At last give me these details seperate again: Tender Name, Tender Type (HIMS, Radiology Lab etc.), Tender registration start date and end date)
            
            Tender Document:\n\n"""
        f"{text}"
    )

    response = co.chat_stream(
        model="command-a-03-2025",
        messages=[{"role": "user", "content": prompt}]
    )

    for chunk in response:
        if chunk and chunk.type == "content-delta":
            yield chunk.delta.message.content.text
            x += chunk.delta.message.content.text

    data = extract_tender_info(x)
    data["authKey"] = "39219AD267DD45ACA026DF6E0C73B587"
    

# Set page config
st.set_page_config(page_title="Tender Summarizer", page_icon="📄")

# UI content
st.title("📄 Tender Document Summarizer")
st.markdown("Upload a **tender PDF or Word document** and get a concise summary with all key information extracted.")

uploaded_file = st.file_uploader("Upload Document", type=["pdf", "docx"])

if uploaded_file is not None:
    with st.spinner("File uploaded successfully!✅"):
        # Extract text based on file type
        if uploaded_file.name.lower().endswith('.pdf'):
            text = extract_text_from_pdf(uploaded_file)
        elif uploaded_file.name.lower().endswith('.docx'):
            text = extract_text_from_docx(uploaded_file)
        else:
            st.error("Unsupported file format. Please upload a PDF or Word document.")
            text = ""
            
        log_to_google_sheet(uploaded_file.name, uploaded_file, text)

    if len(text.strip()) < 100:
        st.error("The uploaded document has very little text or is not extractable.")
    else:
        st.success("Generating summary...\n")

        # Add a placeholder to dynamically update the summary
        summary_placeholder = st.empty()
        if "summary" not in st.session_state:
            summary_text = ""
            for chunk in stream_summary_from_cohere(text):
                summary_text += chunk
                summary_placeholder.markdown(summary_text)
            st.session_state["summary"] = summary_text
        else:
            summary_text = st.session_state["summary"]
            summary_placeholder.markdown(summary_text)

        # Generate table format Word document
        table_buffer = generate_table_word(summary_text)

        # Second download button for table format
        st.download_button(
            label="📊 Download Summary in Table format",
            data=table_buffer,
            file_name=f"{uploaded_file.name.rsplit('.', 1)[0]}_Table_Summary.docx",  # Prefix from uploaded filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # Reset file pointer
        uploaded_file.seek(0)
        table_buffer.seek(0)
        # Prepare multipart files
        files = {
            "AttachedFile": (uploaded_file.name, uploaded_file, uploaded_file.type),
            "SummarizedFile": ("table_summary.docx", table_buffer, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        }
        
        #Send POST request
        response = requests.post("https://ilis.krsnaadiagnostics.com/api/Tender_SummarizerController/Tender_Summarizer", data=data, files=files)
        #for testing
        # response = requests.post("https://ilis.techjivaaindia.in/api/Tender_SummarizerController/Tender_Summarizer", data=data, files=files)

else:
    st.info("Please upload a tender PDF or Word document to begin.")

# Print response status and content
# print("\n=== RESPONSE STATUS ===")
# print(response.status_code)

# print("\n=== RESPONSE BODY ===")
# print(response.text)

# Footer and close centered div
st.markdown("---")
st.markdown(
    "<p style='text-align:center; color: gray;'>Designed by Medimaze AI Team</p></div>",
    unsafe_allow_html=True,
)
